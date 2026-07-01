#!/usr/bin/env python3
"""
Skrypt generujacy plik DOCX z wnioskiem o skierowanie pytan prejudycjalnych do TSUE.
Wczytuje tresc z pliku MD i konwertuje na dokument DOCX z poprawnym formatowaniem.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


def create_document():
    """Create document with proper margins and default style."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5

    return doc


def set_run_font(run, size=12, bold=False, italic=False):
    """Set font properties on a run."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_formatted_text(para, text, bold=False, italic=False, size=12):
    """Add text with inline **bold** markers parsed."""
    if '**' in text:
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                set_run_font(run, size=size, bold=True, italic=italic)
            else:
                run = para.add_run(part)
                set_run_font(run, size=size, bold=bold, italic=italic)
    else:
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)


def add_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


def is_continuation_line(line, current_text=''):
    """Check if a line should be merged with the previous one.
    Short address-like lines are NOT merged."""
    stripped = line.strip()
    if stripped == '':
        return False
    if stripped.startswith('#'):
        return False
    if stripped.startswith('- '):
        return False
    if stripped.startswith('> '):
        return False
    if stripped.startswith('|'):
        return False
    if stripped == '---':
        return False
    if re.match(r'^\([a-z]\)\s', stripped):
        return False
    if re.match(r'^\d+\.\s', stripped):
        return False
    if stripped.startswith('**'):
        return False
    # Don't merge short address/header lines (less than 60 chars each)
    if len(current_text) < 60 and len(stripped) < 60:
        return False
    return True


def parse_md_and_build_docx(md_path, output_path):
    """Parse the markdown file and build the DOCX document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = create_document()

    # Track if we've already processed the signature (to avoid duplication)
    signature_started = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # H1 heading (title)
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            run = para.add_run(text)
            set_run_font(run, size=14, bold=True)
            i += 1
            continue

        # H2 heading
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            if text == 'WNIOSEK':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                set_run_font(run, size=14, bold=True)
            elif text.startswith('o skierowanie'):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                set_run_font(run, size=12, bold=False)
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(text)
                set_run_font(run, size=14, bold=True)
            i += 1
            continue

        # H3 heading
        if line.startswith('### '):
            text = line[4:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(4)
            add_formatted_text(para, text, bold=True, size=12)
            i += 1
            continue

        # H4 heading
        if line.startswith('#### '):
            text = line[5:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(text)
            set_run_font(run, size=12, bold=True)
            i += 1
            continue

        # Block quote
        if line.startswith('> '):
            quote_text = line[2:].strip()
            # Collect multi-line quotes
            while i + 1 < len(lines) and lines[i + 1].rstrip('\n').startswith('> '):
                i += 1
                quote_text += ' ' + lines[i].rstrip('\n')[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.5)
            para.paragraph_format.line_spacing = 1.5
            # Remove surrounding quotation marks for display
            display_text = quote_text.strip('"').strip('\u201e').strip('\u201d')
            run = para.add_run('\u201e' + display_text + '\u201d')
            set_run_font(run, size=12, italic=True)
            i += 1
            continue

        # Table
        if line.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
            rows = []
            for tl in table_lines[2:]:
                row = [cell.strip() for cell in tl.split('|')[1:-1]]
                rows.append(row)

            num_cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ci, header in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(header)
                set_run_font(run, size=11, bold=True)

            for ri, row_data in enumerate(rows):
                for ci in range(min(num_cols, len(row_data))):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(row_data[ci])
                    set_run_font(run, size=11)

            add_table_borders(table)
            doc.add_paragraph()
            continue

        # Bullet point
        if line.startswith('- '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = 1.5
            add_formatted_text(para, '\u2013 ' + text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            # Collect continuation lines
            while i + 1 < len(lines) and is_continuation_line(lines[i + 1].rstrip('\n'), text):
                i += 1
                text += ' ' + lines[i].rstrip('\n').strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = 1.5
            add_formatted_text(para, f'{num}. {text}')
            i += 1
            continue

        # Lettered list (a), (b), (c), (d)
        letter_match = re.match(r'^\(([a-z])\)\s+(.+)', line)
        if letter_match:
            letter = letter_match.group(1)
            text = letter_match.group(2)
            while i + 1 < len(lines) and is_continuation_line(lines[i + 1].rstrip('\n'), text):
                i += 1
                text += ' ' + lines[i].rstrip('\n').strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.5)
            para.paragraph_format.first_line_indent = Cm(-0.5)
            para.paragraph_format.line_spacing = 1.5
            add_formatted_text(para, f'({letter}) {text}')
            i += 1
            continue

        # Signature section - detect and handle specially
        if line.strip().startswith('**Podpis wnioskodawcy:**'):
            signature_started = True
            i += 1
            # Skip remaining MD signature lines
            while i < len(lines):
                i += 1
            break

        # Regular paragraph - lines starting with ** are separate paragraphs
        text = line.strip()
        # Only collect continuation lines if the current line does NOT start with **
        if not text.startswith('**'):
            while i + 1 < len(lines) and is_continuation_line(lines[i + 1].rstrip('\n'), text):
                i += 1
                text += ' ' + lines[i].rstrip('\n').strip()

        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5

        # Right-align the date
        if 'dnia 30 czerwca 2026' in text:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        add_formatted_text(para, text)
        i += 1

    # Add signature block
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run('Podpis wnioskodawcy:')
    set_run_font(run, size=12, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run('.......................................')
    set_run_font(run, size=12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run('Tomasz \u015awiestowski')
    set_run_font(run, size=12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run('ul. Czarnole\u015bna 43C/3')
    set_run_font(run, size=12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run('41-709 Ruda \u015al\u0105ska')
    set_run_font(run, size=12)

    doc.save(output_path)
    print(f'Dokument zapisany: {output_path}')
    print(f'Liczba akapitow: {len(doc.paragraphs)}')
    print(f'Liczba tabel: {len(doc.tables)}')


if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, '5_wniosek_pytania_prejudycjalne_WSA.md')
    output_path = os.path.join(script_dir, '5_wniosek_pytania_prejudycjalne_WSA.docx')
    parse_md_and_build_docx(md_path, output_path)
