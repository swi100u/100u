#!/usr/bin/env python3
"""
Skrypt generujacy dokument .docx z wnioskiem o wylaczenie sedziego,
pytania prejudycjalne do TSUE i zagadnienie prawne do SN.

Wymaga: pip install python-docx
Uzycie: python3 generuj_wniosek.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime


def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_paragraph_bold(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def add_normal(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_numbered(doc, number, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. ")
    run.bold = True
    run2 = p.add_run(text)
    run2.bold = bold
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    return p


def main():
    doc = Document()

    # Ustawienia marginesow
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ===== NAGLOWEK =====
    title = doc.add_heading("WNIOSEK", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "o wylaczenie sedziego, skierowanie pytania prejudycjalnego "
        "do Trybunalu Sprawiedliwosci Unii Europejskiej\n"
        "oraz przedstawienie zagadnienia prawnego Sadowi Najwyzszemu"
    )
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Dane sprawy
    add_normal(doc, "Sąd: Sąd Okręgowy w Gliwicach", bold=True)
    add_normal(doc, "Wydział: VI Wydział Karny Odwoławczy", bold=True)
    add_normal(doc, "Sygn. akt: VI Kz 289/26", bold=True)
    add_normal(doc, "Wnioskodawca: Tomasz Świestowski", bold=True)
    add_normal(doc, "ul. Czarnoleśna 43C/3, 41-709 Ruda Śląska")

    doc.add_paragraph()

    # ===== CZESC I =====
    add_heading_custom(doc, "I. WNIOSEK O WYLACZENIE SEDZIEGO", level=1)

    add_normal(doc, (
        "Na podstawie art. 41 par. 1 k.p.k. wnoszę o wyłączenie od udziału "
        "w niniejszej sprawie:"
    ))

    add_paragraph_bold(doc, "SSR del. Piotr Pawlik")

    add_normal(doc, "od rozpoznawania sprawy o sygn. akt VI Kz 289/26.")

    doc.add_paragraph()
    add_heading_custom(doc, "Uzasadnienie", level=2)

    add_numbered(doc, 1, (
        "W składzie orzekającym w niniejszej sprawie zasiada sędzia sądu rejonowego "
        "delegowany do orzekania w sądzie okręgowym przez Ministra Sprawiedliwości, "
        "pełniącego jednocześnie funkcję Prokuratora Generalnego."
    ))

    add_numbered(doc, 2, (
        "Instytucja delegacji, ukształtowana w art. 77 ustawy — Prawo o ustroju "
        "sądów powszechnych, budzi poważne wątpliwości co do spełnienia standardu "
        "niezależności i bezstronności sądu w rozumieniu:"
    ))

    bullets = [
        "art. 45 ust. 1 Konstytucji RP (prawo do rozpoznania sprawy przez niezależny, bezstronny i niezawisły sąd),",
        "art. 6 ust. 1 Europejskiej Konwencji Praw Człowieka,",
        "art. 19 ust. 1 Traktatu o Unii Europejskiej,",
        "art. 47 Karty Praw Podstawowych UE."
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

    add_numbered(doc, 3, "W szczególności wskazać należy, że:")

    specifics = [
        "delegacja następuje na podstawie jednostronnej decyzji organu władzy wykonawczej (Ministra Sprawiedliwości),",
        "może zostać odwołana w każdym czasie, bez podania przyczyny,",
        "brak jest skutecznego środka kontroli sądowej tej decyzji,",
        "system ten tworzy obiektywną zależność organizacyjno-funkcjonalną sędziego od organu władzy wykonawczej, który jednocześnie pełni funkcję Prokuratora Generalnego — a więc strony w wielu sprawach karnych."
    ]
    for s in specifics:
        doc.add_paragraph(s, style='List Bullet')

    add_numbered(doc, 4, (
        "Zagadnienie to jest przedmiotem aktualnie zawisłego przed Sądem Najwyższym "
        "pytania prawnego (sygn. I KZP 1/26), skierowanego przez Sąd Okręgowy w Rzeszowie, "
        "dotyczącego kwestii, czy udział sędziego delegowanego przez Ministra Sprawiedliwości "
        "stanowi nienależytą obsadę sądu w rozumieniu art. 439 par. 1 pkt 2 k.p.k."
    ))

    add_numbered(doc, 5, (
        "Niniejsza sprawa dotyczy kwestii o szczególnym ciężarze gatunkowym — rozstrzyga "
        "o pozbawieniu wolności wnioskodawcy w drodze obserwacji psychiatrycznej. Zgodnie z wyrokiem "
        "Europejskiego Trybunału Praw Człowieka z dnia 14.10.2021 r. w sprawie M.B. przeciwko Polsce "
        "(skarga nr 60157/15), pozbawienie wolności osoby z zaburzeniami psychicznymi wymaga "
        "spełnienia szczególnie rygorystycznych standardów legalności."
    ))

    add_numbered(doc, 6, (
        "W sytuacji, gdy o pozbawieniu wolności decyduje sędzia, którego niezależność strukturalna "
        "budzi obiektywne wątpliwości, zachodzi uzasadniona obawa, że standard art. 5 ust. 1 lit. e EKPC "
        "oraz art. 6 ust. 1 EKPC nie zostanie dochowany."
    ))

    # ===== CZESC II =====
    doc.add_page_break()
    add_heading_custom(doc, "II. ARGUMENT Z AKTUALNOŚCI OPINII BIEGŁYCH", level=1)
    add_heading_custom(doc, "(Wyrok ETPCz M.B. przeciwko Polsce, skarga nr 60157/15)", level=2)

    add_heading_custom(doc, "A. Stan faktyczny", level=3)

    add_numbered(doc, 7, (
        "W sprawie 4096-6.Ds.94.2025 biegli powołani przez prokuraturę twierdzą, że "
        "przeprowadzenie badania sądowo-psychiatrycznego w trybie ambulatoryjnym NIE JEST MOŻLIWE "
        "i konieczna jest obserwacja w zakładzie psychiatrycznym w Toszku."
    ))

    add_numbered(doc, 8, (
        "Jednocześnie istnieje nowsza opinia sądowo-psychiatryczna, wydana po przeprowadzeniu "
        "badania w trybie ambulatoryjnym — co bezpośrednio przeczy twierdzeniu biegłych "
        "o niemożności takiego badania."
    ))

    add_numbered(doc, 9, (
        "Sprzeczność ta nie została wyjaśniona przez sąd ani prokuratora przed podjęciem "
        "decyzji o skierowaniu na obserwację."
    ))

    add_heading_custom(doc, "B. Standard wynikający z wyroku ETPCz M.B. p. Polsce", level=3)

    add_numbered(doc, 10, (
        "W wyroku z dnia 14.10.2021 r. Europejski Trybunał Praw Człowieka stwierdził "
        "naruszenie art. 5 ust. 1 lit. e Konwencji i wskazał trzy minimalne warunki "
        "legalności pozbawienia wolności osoby z zaburzeniami psychicznymi:"
    ))

    conditions = [
        "(a) Występowanie zaburzeń psychicznych musi zostać ustalone w rzetelny sposób, przez kompetentne osoby, w oparciu o obiektywną i aktualną wiedzę medyczną.",
        "(b) Zaburzenie psychiczne musi być tego rodzaju i takiej głębokości, że wymaga przymusowego leczenia szpitalnego (lub obserwacji).",
        "(c) Zasadność dalszego pobytu musi być uzależniona od utrzymywania się zaburzeń psychicznych."
    ]
    for c in conditions:
        doc.add_paragraph(c, style='List Bullet')

    add_numbered(doc, 11, "Trybunał podkreślił, że:")

    add_quote(doc, (
        "\"Ocena medyczna musi opierać się na rzeczywistym stanie zdrowia psychicznego "
        "danej osoby, a nie wyłącznie na gebeurtenościach z przeszłości.\""
    ))

    add_quote(doc, (
        "\"Opinia lekarska nie jest wystarczająca dla uzasadnienia pozbawienia wolności, "
        "jeśli od jej sporządzenia upłynął znaczny czas.\""
    ))

    add_normal(doc, (
        "Sąd ma obowiązek wyjaśnić sprzeczności między opiniami biegłych a innymi "
        "dowodami medycznymi PRZED pozbawieniem wolności."
    ))

    add_numbered(doc, 12, (
        "W sprawie M.B. Trybunał uznał opinie z 17.01.2014 r. i 3.02.2014 r. za "
        "niewystarczająco aktualne w momencie umieszczenia skarżącego w szpitalu "
        "(4.08.2015 r.) — upływ ok. 18 miesięcy, przy jednoczesnym ignorowaniu dowodów "
        "poprawy stanu zdrowia."
    ))

    add_heading_custom(doc, "C. Zastosowanie do niniejszej sprawy", level=3)

    add_numbered(doc, 13, "Analogia jest bezpośrednia:")

    analogies = [
        "Istnieje nowsza opinia wydana po badaniu ambulatoryjnym, która przeczy twierdzeniom biegłych o niemożności takiego badania.",
        "Sąd nie wyjaśnił sprzeczności między tą opinią a wnioskiem o obserwację.",
        "Skierowanie na obserwację w Toszku — stanowiące de facto pozbawienie wolności — opiera się na twierdzeniach biegłych, które zostały podważone przez późniejsze badanie ambulatoryjne.",
        "Zgodnie ze standardem M.B. p. Polsce, sąd POWINIEN BYŁ oprzeć się na najaktualniejszej opinii medycznej lub co najmniej wyjaśnić rozbieżność przed pozbawieniem wolności."
    ]
    for a in analogies:
        doc.add_paragraph(a, style='List Bullet')

    add_numbered(doc, 14, (
        "Odmowa uwzględnienia nowszej opinii i oparcie decyzji o pozbawieniu wolności "
        "wyłącznie na opinii biegłych kwestionujących możliwość badania ambulatoryjnego "
        "— gdy takie badanie faktycznie przeprowadzono — stanowi naruszenie art. 5 ust. 1 "
        "lit. e EKPC w świetle wyroku M.B. p. Polsce."
    ))

    # ===== CZESC III =====
    doc.add_page_break()
    add_heading_custom(doc, "III. WNIOSEK O SKIEROWANIE PYTANIA PREJUDYCJALNEGO DO TSUE", level=1)

    add_normal(doc, (
        "Na podstawie art. 267 Traktatu o funkcjonowaniu Unii Europejskiej wnoszę "
        "o skierowanie do Trybunału Sprawiedliwości Unii Europejskiej następującego "
        "pytania prejudycjalnego:"
    ))

    add_heading_custom(doc, "Pytanie:", level=3)

    add_quote(doc, (
        "Czy art. 19 ust. 1 Traktatu o Unii Europejskiej w związku z art. 47 Karty "
        "Praw Podstawowych Unii Europejskiej należy interpretować w ten sposób, że "
        "sprzeczny z wymogiem niezależnego sądu ustanowionego ustawą jest system krajowy, "
        "w którym:"
    ))
    add_quote(doc, (
        "(a) sędzia może być delegowany do orzekania w sądzie wyższej instancji "
        "wyłącznie na podstawie decyzji organu władzy wykonawczej (Ministra Sprawiedliwości),"
    ))
    add_quote(doc, (
        "(b) delegacja ta może być odwołana w dowolnym czasie bez podania przyczyny,"
    ))
    add_quote(doc, (
        "(c) brak jest skutecznej kontroli sądowej decyzji o odwołaniu z delegacji,"
    ))
    add_quote(doc, (
        "(d) organ delegujący pełni jednocześnie funkcję Prokuratora Generalnego,"
    ))
    add_quote(doc, (
        "a udział tak delegowanego sędziego w składzie orzekającym w sprawie dotyczącej "
        "pozbawienia wolności (obserwacji psychiatrycznej) skutkuje naruszeniem prawa "
        "do sądu ustanowionego ustawą?"
    ))

    add_heading_custom(doc, "Uzasadnienie obowiązku skierowania pytania", level=3)

    add_numbered(doc, 15, (
        "Na podstawie art. 267 akapit 3 TFUE sąd krajowy, którego orzeczenia nie podlegają "
        "zaskarżeniu w ramach krajowego toku instancji, MA OBOWIĄZEK skierowania pytania "
        "prejudycjalnego do TSUE, jeżeli powstaje wątpliwość co do wykładni prawa Unii Europejskiej."
    ))

    add_numbered(doc, 16, (
        "Zgodnie z orzecznictwem Trybunału Sprawiedliwości UE, w szczególności wyrokiem "
        "w sprawie C-283/81 CILFIT, odstąpienie od skierowania pytania prejudycjalnego jest "
        "dopuszczalne wyłącznie w sytuacji, gdy prawidłowe stosowanie prawa Unii jest tak "
        "oczywiste, że nie pozostawia żadnych racjonalnych wątpliwości interpretacyjnych "
        "(acte clair)."
    ))

    add_numbered(doc, 17, (
        "W niniejszej sprawie wątpliwość co do zgodności krajowego modelu delegacji "
        "sędziowskiej z prawem UE NIE JEST OCZYWISTA (nie zachodzi acte clair), o czym świadczy:"
    ))

    evidence = [
        "toczące się postępowanie przed Sądem Najwyższym (I KZP 1/26),",
        "rozbieżne orzecznictwo krajowe,",
        "dotychczasowe orzecznictwo TSUE (C-487/19, C-791/19, C-824/18 A.B., C-506/04 Wilson, C-585/18, C-624/18, C-625/18)."
    ]
    for e in evidence:
        doc.add_paragraph(e, style='List Bullet')

    add_numbered(doc, 18, (
        "Nawet jeśli sąd rozpoznający niniejszą sprawę nie jest formalnie sądem \"ostatniej instancji\", "
        "to w świetle wyroku TSUE w sprawie C-561/19 Consorzio Italian Management każdy sąd krajowy "
        "MOŻE zadać pytanie prejudycjalne, a w sytuacji istotnej wątpliwości co do prawa UE powinien "
        "rozważyć takie skierowanie. Odmowa wymaga pełnego uzasadnienia w świetle standardu CILFIT."
    ))

    add_numbered(doc, 19, (
        "Ponadto, w sytuacji gdy przedmiotem sprawy jest POZBAWIENIE WOLNOŚCI (obserwacja psychiatryczna), "
        "standard ochrony praw jednostki wymaga szczególnej staranności w zapewnieniu niezależności "
        "sądu orzekającego, co dodatkowo uzasadnia skierowanie pytania prejudycjalnego."
    ))

    # ===== CZESC IV =====
    doc.add_page_break()
    add_heading_custom(doc, "IV. WNIOSEK O PRZEDSTAWIENIE ZAGADNIENIA PRAWNEGO SĄDOWI NAJWYŻSZEMU", level=1)

    add_normal(doc, (
        "Na podstawie art. 441 par. 1 k.p.k. wnoszę o przedstawienie Sądowi Najwyższemu "
        "do rozstrzygnięcia następującego zagadnienia prawnego:"
    ))

    add_heading_custom(doc, "Zagadnienie:", level=3)

    add_quote(doc, (
        "Czy udział w składzie orzekającym sędziego sądu rejonowego delegowanego do orzekania "
        "w sądzie okręgowym przez Ministra Sprawiedliwości, w trybie art. 77 ustawy — Prawo "
        "o ustroju sądów powszechnych, stanowi nienależytą obsadę sądu w rozumieniu "
        "art. 439 par. 1 pkt 2 k.p.k., a w konsekwencji bezwzględną przyczynę odwoławczą "
        "skutkującą uchyleniem orzeczenia — w sytuacji gdy:"
    ))
    add_quote(doc, "(a) system delegacji przewiduje możliwość swobodnego odwołania sędziego z delegacji przez organ władzy wykonawczej,")
    add_quote(doc, "(b) nie zapewnia skutecznej kontroli sądowej decyzji Ministra Sprawiedliwości,")
    add_quote(doc, "(c) organ delegujący pełni jednocześnie funkcję Prokuratora Generalnego,")
    add_quote(doc, "(d) sprawa dotyczy pozbawienia wolności (zarządzenia obserwacji psychiatrycznej)?")

    add_heading_custom(doc, "Uzasadnienie", level=3)

    add_numbered(doc, 20, "Zagadnienie to wymaga rozstrzygnięcia przez Sąd Najwyższy z uwagi na:")

    reasons = [
        "aktualnie zawisłą przed SN sprawę I KZP 1/26 (pytanie prawne SO Rzeszów),",
        "rozbieżność w orzecznictwie między linią III CZP 103/15 (dopuszczalność delegacji) a nowszymi orzeczeniami kwestionującymi niezależność sędziów delegowanych,",
        "konieczność ujednolicenia praktyki sądowej w sprawach dotyczących pozbawienia wolności."
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    # ===== CZESC V =====
    doc.add_page_break()
    add_heading_custom(doc, "V. WNIOSKI KOŃCOWE", level=1)

    add_normal(doc, "Na podstawie powyższego wnoszę o:")

    final_requests = [
        "Wyłączenie SSR del. Piotra Pawlika od rozpoznania sprawy VI Kz 289/26, z uwagi na uzasadnione wątpliwości co do jego niezależności strukturalnej wynikającej z delegacji przez Ministra Sprawiedliwości.",
        "Skierowanie pytania prejudycjalnego do TSUE w zakresie zgodności systemu delegacji sędziowskiej z art. 19 ust. 1 TUE i art. 47 KPP UE.",
        "Przedstawienie zagadnienia prawnego Sądowi Najwyższemu w zakresie kwalifikacji udziału sędziego delegowanego jako nienależytej obsady sądu (art. 439 par. 1 pkt 2 k.p.k.).",
        "Uwzględnienie standardu aktualności opinii biegłych wynikającego z wyroku ETPCz z 14.10.2021 r. w sprawie M.B. przeciwko Polsce (skarga nr 60157/15) i wyjaśnienie sprzeczności między opiniami biegłych a dowodami ambulatoryjnymi PRZED pozbawieniem wolności."
    ]
    for i, req in enumerate(final_requests, 1):
        doc.add_paragraph(f"{i}. {req}", style='List Number')

    doc.add_paragraph()
    doc.add_paragraph()

    # Podpis
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[miejscowość], dnia [data]")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[podpis wnioskodawcy / obrońcy]")

    # ===== TABELA PODSTAW PRAWNYCH =====
    doc.add_page_break()
    add_heading_custom(doc, "PODSTAWY PRAWNE (zestawienie)", level=2)

    table_data = [
        ("Przepis", "Treść / znaczenie"),
        ("Art. 41 par. 1 k.p.k.", "Wyłączenie sędziego na wniosek strony"),
        ("Art. 439 par. 1 pkt 2 k.p.k.", "Bezwzględna przyczyna odwoławcza — nienależyta obsada sądu"),
        ("Art. 441 par. 1 k.p.k.", "Przedstawienie zagadnienia prawnego Sądowi Najwyższemu"),
        ("Art. 77 u.s.p.", "Delegacja sędziego przez Ministra Sprawiedliwości"),
        ("Art. 267 TFUE", "Pytanie prejudycjalne do TSUE"),
        ("Art. 19 ust. 1 TUE", "Obowiązek zapewnienia skutecznej ochrony sądowej"),
        ("Art. 47 KPP UE", "Prawo do niezależnego i bezstronnego sądu"),
        ("Art. 5 ust. 1 lit. e EKPC", "Legalność pozbawienia wolności osoby z zaburzeniami psychicznymi"),
        ("Art. 6 ust. 1 EKPC", "Prawo do rzetelnego procesu"),
        ("Art. 45 ust. 1 Konstytucji RP", "Prawo do sądu"),
    ]

    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (col1, col2) in enumerate(table_data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # ===== TABELA ORZECZNICTWA =====
    add_heading_custom(doc, "ORZECZNICTWO PRZYWOŁANE", level=2)

    case_data = [
        ("Sygnatura", "Organ", "Znaczenie"),
        ("M.B. p. Polsce, 60157/15", "ETPCz (14.10.2021)", "Aktualność opinii biegłych jako warunek legalności pozbawienia wolności"),
        ("I KZP 1/26", "SN", "Pytanie prawne — delegacja MS a art. 439 par. 1 pkt 2 k.p.k."),
        ("III CZP 103/15", "SN (2016)", "Delegacja sędziego — dopuszczalność co do zasady"),
        ("C-283/81 CILFIT", "TSUE", "Standard acte clair — wyjątek od obowiązku pytania prejudycjalnego"),
        ("C-487/19", "TSUE", "Niezależność sądu — test strukturalny"),
        ("C-791/19", "TSUE", "Systemowe naruszenia niezależności sądów w Polsce"),
        ("C-824/18 A.B.", "TSUE", "Niezależność sądu — powołania sędziowskie"),
        ("C-506/04 Wilson", "TSUE", "Definicja niezależnego sądu"),
        ("C-585/18, C-624/18, C-625/18", "TSUE", "Test niezależności — Izba Dyscyplinarna SN"),
        ("C-561/19 Consorzio", "TSUE", "Zakres obowiązku pytania prejudycjalnego"),
        ("II KK 140/21", "SN", "Skład sądu z udziałem sędziego delegowanego"),
        ("Winterwerp p. Holandii", "ETPCz", "Trzy warunki legalności pozbawienia wolności"),
    ]

    table2 = doc.add_table(rows=len(case_data), cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(case_data):
        row = table2.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Zapis
    filename = "Wniosek_wylaczenie_TSUE_SN_II_Kp_550_25.docx"
    doc.save(filename)
    print(f"Dokument zapisany jako: {filename}")
    print(f"Data generowania: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")


if __name__ == "__main__":
    main()
